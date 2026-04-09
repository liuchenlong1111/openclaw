#!/usr/bin/env python3
"""
Extract text content from Word documents (.docx).
"""

import sys
import argparse

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

def extract_docx(file_path):
    """Extract text from Word document"""
    try:
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting DOCX: {e}", file=sys.stderr)
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(description='Extract text from Word documents')
    parser.add_argument('file_path', help='Path to .docx file')
    parser.add_argument('--output', '-o', help='Output file (default: stdout)')
    args = parser.parse_args()
    
    text = extract_docx(args.file_path)
    
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Extracted text saved to {args.output}", file=sys.stderr)
    else:
        print(text)

if __name__ == '__main__':
    main()
